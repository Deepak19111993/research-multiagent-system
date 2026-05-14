import io
import markdown
from fpdf import FPDF
from docx import Document

def create_markdown_file(content: str) -> bytes:
    """Returns the raw markdown content as bytes."""
    return content.encode('utf-8')

def create_html_file(markdown_content: str) -> bytes:
    """Converts markdown to HTML and returns as bytes."""
    html = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
    
    # Wrap in basic HTML structure for better styling
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }}
            h1, h2, h3 {{ color: #2c3e50; }}
            code {{ background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
            pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            blockquote {{ border-left: 4px solid #ccc; margin: 0; padding-left: 15px; color: #666; }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    return full_html.encode('utf-8')

def create_pdf_file(markdown_content: str) -> bytes:
    """Converts markdown to PDF and returns as bytes."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    
    def sanitize_text(text: str) -> str:
        """Replace common unicode characters with ASCII equivalents for basic FPDF fonts."""
        replacements = {
            '—': '-', '–': '-',
            '“': '"', '”': '"',
            '‘': "'", '’': "'",
            '…': '...',
            '•': '*',
            '™': '(tm)',
            '©': '(c)',
            '®': '(r)'
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Encode to latin-1 and replace unmappable chars with '?' to prevent crashing
        return text.encode('latin-1', 'replace').decode('latin-1')
    
    # Very basic markdown handling for PDF to avoid complex HTML parsing dependencies
    lines = markdown_content.split('\\n')
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
            
        if line.startswith('# '):
            pdf.set_font("helvetica", style='B', size=16)
            pdf.multi_cell(0, 10, sanitize_text(line[2:]))
            pdf.set_font("helvetica", size=12)
        elif line.startswith('## '):
            pdf.set_font("helvetica", style='B', size=14)
            pdf.multi_cell(0, 8, sanitize_text(line[3:]))
            pdf.set_font("helvetica", size=12)
        elif line.startswith('### '):
            pdf.set_font("helvetica", style='B', size=13)
            pdf.multi_cell(0, 7, sanitize_text(line[4:]))
            pdf.set_font("helvetica", size=12)
        else:
            # Strip markdown bold/italic tags for simplicity in basic PDF
            clean_line = line.replace('**', '').replace('*', '')
            pdf.multi_cell(0, 6, sanitize_text(clean_line))
            
    return bytes(pdf.output(dest='S'))

def create_word_file(markdown_content: str) -> bytes:
    """Creates a Word document from markdown text (basic formatting) and returns as bytes."""
    doc = Document()
    
    # A very basic parser: split by double newlines to create paragraphs
    paragraphs = markdown_content.split('\\n\\n')
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
            
        if p.startswith('# '):
            doc.add_heading(p[2:], level=1)
        elif p.startswith('## '):
            doc.add_heading(p[3:], level=2)
        elif p.startswith('### '):
            doc.add_heading(p[4:], level=3)
        elif p.startswith('- ') or p.startswith('* '):
            doc.add_paragraph(p, style='List Bullet')
        else:
            doc.add_paragraph(p)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
