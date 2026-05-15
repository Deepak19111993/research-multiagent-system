import io
import markdown
from fpdf import FPDF
from docx import Document

def create_markdown_file(content: str) -> bytes:
    """Returns the raw markdown content as bytes."""
    return content.encode('utf-8')

def create_html_file(markdown_content: str) -> bytes:
    """Converts markdown to HTML and returns as bytes."""
    html = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
    
    # Wrap in basic HTML structure for better styling
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }}
            h1, h2, h3 {{ color: #2c3e50; }}
            code {{ background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
            pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
            blockquote {{ border-left: 4px solid #ccc; margin: 0; padding-left: 15px; color: #666; }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    return full_html.encode('utf-8')

def create_pdf_file(markdown_content: str) -> bytes:
    """Converts markdown to PDF with Unicode safety and robust error handling."""
    import markdown
    import re
    from fpdf import FPDF
    
    def clean_special_chars(text):
        replacements = {
            '\u2014': '-', '\u2013': '-',
            '\u201c': '"', '\u201d': '"',
            '\u2018': "'", '\u2019': "'",
            '\u2026': '...', '\u2022': '*',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    # Strip images and fix special Unicode characters (em-dash, etc.)
    clean_markdown = re.sub(r'!\[.*?\]\(.*?\)', '', markdown_content)
    clean_markdown = clean_special_chars(clean_markdown)
    
    # Convert markdown to simple HTML
    html = markdown.markdown(clean_markdown, extensions=['extra'])
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    
    try:
        pdf.write_html(html)
    except Exception:
        # Fallback: start a fresh page and write plain text if HTML rendering fails
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        lines = clean_markdown.split('\n')
        for line in lines:
            # Encode with latin-1 and replace unmappable chars with ? to avoid crash
            safe_line = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, safe_line)
        
    return bytes(pdf.output())

def create_word_file(markdown_content: str) -> bytes:
    """Creates a Word document with improved formatting preservation."""
    import re
    doc = Document()
    
    # Split by double newlines for paragraphs
    paragraphs = markdown_content.split('\n\n')
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
            
        if p.startswith('#'):
            level = 0
            while level < len(p) and p[level] == '#':
                level += 1
            title = p[level:].strip()
            if title:
                doc.add_heading(title, level=min(level, 4))
        elif p.startswith('- ') or p.startswith('* '):
            doc.add_paragraph(p[2:].strip(), style='List Bullet')
        else:
            # For Word, we can't easily import HTML, so we do basic inline formatting replacement
            # Handle Bold: **text**
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', p)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
